# apps/quiz/services.py
"""
Service layer for the AI Quiz Generator.

Responsibilities:
  1. Download a study material file from a URL (Cloudinary in production).
  2. Extract its text — PDF (pypdf) or DOCX (python-docx).
  3. Ask OpenAI to generate quiz questions in a strict JSON schema.
  4. Parse, validate, and return the questions.

The frontend never has to send file bytes — it just passes a SavedMaterial
id, and the backend resolves the file_url and does everything else.
"""
import io
import json
import logging
import re
import uuid
from typing import Any

import requests
from django.conf import settings

from apps.ai import ai_router  # Phase 3: quiz generation routes through the quiz lane

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────
# 1.  File download
# ─────────────────────────────────────────────────────────────

class FileFetchError(Exception):
    """Raised when we can't pull the source file."""


def _sign_cloudinary(url: str) -> str:
    """Return a signed Cloudinary URL so raw / PDF assets download even when
    public PDF & ZIP delivery is disabled on the account. Signed URLs are
    Cloudinary's sanctioned way to deliver restricted file types. Falls back
    to the original url on any problem (non-Cloudinary url, bad config, etc)."""
    if not url or "res.cloudinary.com" not in url:
        return url
    try:
        import re as _re
        import cloudinary.utils as _cu
        m = _re.match(r"https?://res\.cloudinary\.com/[^/]+/([^/]+)/([^/]+)/(.+)$", url)
        if not m:
            return url
        rtype, dtype, rest = m.group(1), m.group(2), m.group(3)
        rest = _re.sub(r"^s--[^/]+--/", "", rest)
        rest = _re.sub(r"^v\d+/", "", rest)
        return _cu.private_download_url(rest, "", resource_type=rtype, type=dtype) or url
    except Exception:
        return url


def download_bytes(url: str, timeout: int = 30, max_bytes: int = 25 * 1024 * 1024) -> bytes:
    """Stream a file from a URL and cap the size to avoid memory blowups.

    25 MB cap matches the backend's video upload limit; PDFs and DOCXs
    are usually well under 5 MB.
    """
    try:
        url = _sign_cloudinary(url)
        resp = requests.get(url, stream=True, timeout=timeout)
        resp.raise_for_status()
    except requests.RequestException as e:
        raise FileFetchError(f"Could not fetch file: {e}") from e

    buf = bytearray()
    for chunk in resp.iter_content(chunk_size=64 * 1024):
        if not chunk:
            continue
        buf.extend(chunk)
        if len(buf) > max_bytes:
            raise FileFetchError(
                f"File exceeds the {max_bytes // 1024 // 1024} MB limit.")
    return bytes(buf)


# ─────────────────────────────────────────────────────────────
# 2.  Text extraction
# ─────────────────────────────────────────────────────────────

class ExtractionError(Exception):
    """Raised when we can't pull readable text out of the file."""


def _normalise_whitespace(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ExtractionError(
            "pypdf is not installed. Add `pypdf>=4.0` to requirements.txt."
        ) from e

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")  # try empty password
            except Exception:
                raise ExtractionError("This PDF is encrypted and cannot be read.")
        chunks = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception as e:                    # noqa: BLE001
                logger.warning("PDF page extract failed: %s", e)
        return _normalise_whitespace("\n\n".join(chunks))
    except ExtractionError:
        raise
    except Exception as e:                            # noqa: BLE001
        raise ExtractionError(f"Couldn't read this PDF: {e}") from e


def extract_text_from_docx(data: bytes) -> str:
    try:
        import docx                                  # python-docx
    except ImportError as e:
        raise ExtractionError(
            "python-docx is not installed. Add `python-docx>=1.1` to requirements.txt."
        ) from e

    try:
        document = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in document.paragraphs if p.text]
        # Also pull in table cells — important for study notes in tabular form.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return _normalise_whitespace("\n".join(paragraphs))
    except Exception as e:                            # noqa: BLE001
        raise ExtractionError(f"Couldn't read this DOCX: {e}") from e


def extract_text(file_url: str, file_name: str = "", file_type: str = "") -> str:
    """Pick the right extractor based on extension/MIME and return the text."""
    raw = download_bytes(file_url)

    name = (file_name or file_url).lower()
    ftype = (file_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ftype:
        return extract_text_from_pdf(raw)
    if name.endswith(".docx") or "wordprocessingml" in ftype or "doc" in ftype:
        return extract_text_from_docx(raw)

    raise ExtractionError(
        "Only PDF and DOCX files are supported for quiz generation right now. "
        f"This material looks like: {file_type or 'unknown'}.")


# ─────────────────────────────────────────────────────────────
# 3.  OpenAI quiz generation
# ─────────────────────────────────────────────────────────────

class GenerationError(Exception):
    """Raised when the LLM call fails or returns garbage."""


# Cap the text we send so we stay well within model context.
# ~16k chars ≈ ~4k tokens — comfortable for gpt-4o-mini's 128k window
# and keeps cost predictable.
_MAX_INPUT_CHARS = 16_000

_SYSTEM_PROMPT = """\
You are an expert academic quiz writer. You will be given study material and
must generate quiz questions strictly grounded in that material.

Rules:
  • Only ask about concepts, facts, definitions, or relationships present in
    the material. Do NOT invent content.
  • Each question must be answerable purely from the material.
  • For MCQs: exactly 4 options labelled A, B, C, D — one is unambiguously
    correct, the other three are plausible distractors.
  • For true/false: the answer must be either true or false (boolean).
  • For short answer: keep the expected answer to ≤ 12 words.
  • Always include a brief `explanation` (≤ 30 words) citing the material.
  • Output ONLY a valid JSON object. No prose. No markdown fences.

Output schema:
{
  "questions": [
    {
      "id":             "q1",
      "type":           "mcq" | "true_false" | "short",
      "question":       "...",
      "options":        ["...","...","...","..."],   // mcq only; omit otherwise
      "correct_answer": "B" | true | "expected text",
      "explanation":    "..."
    }
  ]
}
"""


def _user_prompt(text: str, *, num_questions: int, difficulty: str,
                 question_types: list[str], subject: str) -> str:
    types_clause = (
        f"Use only these question types: {', '.join(question_types)}."
        if question_types else
        "Use a mix of mcq, true_false, and short."
    )
    subject_clause = f"Subject context: {subject}." if subject else ""
    truncated = text[:_MAX_INPUT_CHARS]
    if len(text) > _MAX_INPUT_CHARS:
        truncated += "\n\n[...truncated for length...]"

    return (
        f"Generate exactly {num_questions} questions at {difficulty} difficulty.\n"
        f"{types_clause}\n"
        f"{subject_clause}\n\n"
        "── STUDY MATERIAL ──\n"
        f"{truncated}\n"
        "────────────────────\n"
    )


def _call_openai(messages: list[dict], *, model: str) -> tuple[str, int]:
    """Returns (raw_json_string, tokens_used). Uses the official SDK."""
    import os as _os
    api_key = (getattr(settings, "GEMINI_API_KEY", None)
               or _os.environ.get("GEMINI_API_KEY"))
    if not api_key:
        raise GenerationError(
            "GEMINI_API_KEY is not configured on the server.")

    try:
        from openai import OpenAI
    except ImportError as e:
        raise GenerationError(
            "`openai` package not installed. Add `openai>=1.40` to requirements.txt."
        ) from e

    client = OpenAI(
        api_key=api_key,
        base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
    )

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.4,
            response_format={"type": "json_object"},
        )
    except Exception as e:                            # noqa: BLE001
        raise GenerationError(f"AI request failed: {e}") from e

    content = (resp.choices[0].message.content or "").strip()
    tokens  = getattr(resp.usage, "total_tokens", 0) if resp.usage else 0
    return content, tokens


def _strip_fences(raw: str) -> str:
    """Recover JSON when a model wraps it in ```fences``` or adds stray prose."""
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-zA-Z]*\n?", "", raw)
        raw = re.sub(r"\n?```\s*$", "", raw).strip()
    if not raw.startswith("{"):
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            raw = m.group(0)
    return raw


def _validate_questions(raw: str, *, max_q: int) -> list[dict]:
    """Parse + sanitise the JSON we got back from the LLM."""
    try:
        payload = json.loads(_strip_fences(raw))
    except json.JSONDecodeError as e:
        raise GenerationError(f"AI returned malformed JSON: {e}") from e

    questions = payload.get("questions") if isinstance(payload, dict) else None
    if not isinstance(questions, list) or not questions:
        raise GenerationError("AI returned no questions.")

    cleaned: list[dict] = []
    for i, q in enumerate(questions[:max_q], start=1):
        if not isinstance(q, dict):
            continue
        qtype = (q.get("type") or "").lower().strip()
        if qtype not in ("mcq", "true_false", "short"):
            continue
        question_text = (q.get("question") or "").strip()
        if not question_text:
            continue

        item: dict[str, Any] = {
            "id":          q.get("id") or f"q{i}",
            "type":        qtype,
            "question":    question_text,
            "explanation": (q.get("explanation") or "").strip(),
        }

        if qtype == "mcq":
            opts = q.get("options") or []
            if not isinstance(opts, list) or len(opts) != 4:
                continue
            item["options"]        = [str(o).strip() for o in opts]
            item["correct_answer"] = str(q.get("correct_answer") or "").strip().upper()
            if item["correct_answer"] not in ("A", "B", "C", "D"):
                continue
        elif qtype == "true_false":
            ans = q.get("correct_answer")
            if isinstance(ans, str):
                ans = ans.strip().lower() in ("true", "t", "yes", "1")
            item["correct_answer"] = bool(ans)
        else:  # short
            item["correct_answer"] = str(q.get("correct_answer") or "").strip()
            if not item["correct_answer"]:
                continue

        cleaned.append(item)

    if not cleaned:
        raise GenerationError("None of the AI's questions passed validation.")
    return cleaned


def generate_quiz_from_material(
    *,
    file_url: str,
    file_name: str,
    file_type: str,
    subject: str = "",
    num_questions: int = 10,
    difficulty: str = "medium",
    question_types: list[str] | None = None,
) -> dict:
    """Top-level entry point used by the view.

    Returns a dict: {questions, tokens_used, model_name}.
    Raises FileFetchError / ExtractionError / GenerationError on failure.
    """
    num_questions = max(3, min(int(num_questions or 10), 25))
    difficulty = (difficulty or "medium").lower()
    if difficulty not in ("easy", "medium", "hard", "mixed"):
        difficulty = "medium"
    question_types = [t for t in (question_types or [])
                      if t in ("mcq", "true_false", "short")]

    text = extract_text(file_url, file_name=file_name, file_type=file_type)
    if len(text) < 200:
        raise ExtractionError(
            "Not enough readable text in this file to build a meaningful quiz.")

    result = ai_router.complete(
        "quiz",
        messages=[
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user",   "content": _user_prompt(
                text,
                num_questions=num_questions,
                difficulty=difficulty,
                question_types=question_types,
                subject=subject,
            )},
        ],
        max_tokens=2048,
        temperature=0.4,
        response_format={"type": "json_object"},
    )
    raw = result.get("text") or ""
    if not raw:
        raise GenerationError(result.get("error") or "The AI returned an empty response.")

    questions = _validate_questions(raw, max_q=num_questions)
    return {
        "questions":   questions,
        "tokens_used": 0,
        "model_name":  result.get("provider") or "router",
    }


# ─────────────────────────────────────────────────────────────
# 4.  Grading (server-side, so we don't trust the client)
# ─────────────────────────────────────────────────────────────

def _normalise_short(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def grade_attempt(quiz_questions: list[dict], answers: dict) -> tuple[int, int, list[dict]]:
    """Returns (correct_count, total, per_question_breakdown)."""
    breakdown: list[dict] = []
    correct = 0

    for q in quiz_questions:
        qid = q.get("id")
        user_ans = answers.get(qid)
        is_correct = False

        if q["type"] == "mcq":
            is_correct = (str(user_ans or "").strip().upper()
                          == q["correct_answer"])
        elif q["type"] == "true_false":
            if isinstance(user_ans, str):
                user_ans = user_ans.strip().lower() in ("true", "t", "yes", "1")
            is_correct = bool(user_ans) == bool(q["correct_answer"])
        elif q["type"] == "short":
            # Tolerant comparison — exact OR contains the canonical answer.
            ua = _normalise_short(str(user_ans or ""))
            ca = _normalise_short(q["correct_answer"])
            is_correct = bool(ua) and (ua == ca or ca in ua or ua in ca)

        if is_correct:
            correct += 1
        breakdown.append({
            "id":             qid,
            "user_answer":    user_ans,
            "correct_answer": q["correct_answer"],
            "is_correct":     is_correct,
            "explanation":    q.get("explanation", ""),
        })

    return correct, len(quiz_questions), breakdown